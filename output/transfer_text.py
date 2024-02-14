import shutil
from docx import Document
from xml.sax.saxutils import escape
import re





def copy_word_document(original_path, copy_path):
    try:
        shutil.copy2(original_path, copy_path)
        print(f"Copy successful! {original_path} copied to {copy_path}")
    except FileNotFoundError:
        print(f"Error: The original file {original_path} was not found.")
    except PermissionError:
        print(f"Error: Permission denied. Unable to copy {original_path}.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

original_document_path = './data/raw/21-342562.35_Training_Report.docx'
copy_document_path = 'output/21-342562.35_Generated_Report.docx'
copy_word_document(original_document_path, copy_document_path)


def get_section(start_section, end_section):
    # TODO: optimize logic to one loop
    # TODO: track start:stop sections in dict
    file_path = "output/output_text.txt"

    i = 0
    sections = []
    with open(file_path, 'r') as file:
        for line in file:
            i += 1
            if start_section in line:
                sections.append(i)
                print(2, i)
            if end_section in line:
                sections.append(i)
        sections = sections[2:]
        print(sections)
        print()

    section_lines = ""  
    with open(file_path, 'r') as file:
            for i, line in enumerate(file, start=1):
                if sections[0]+1 < i < sections[1]:
                    section_lines += line
    print(section_lines)
    return section_lines

start_section = input("Section to transfer: ")
end_section = input("Next Section: ")
new_text = get_section(start_section, end_section)


# Find line that says 2.4.1
target_string = '2.4.1'
stop_string = '2.4.2'

big_string_for_word = ''

file_path = 'output/output_text.txt'


with open(file_path, 'r') as file:
    use_line = False
    for line in file:
        if target_string in line:
            use_line = True
        if stop_string in line:
           use_line = False
        #print(line)
        if use_line:
            big_string_for_word += line

# print(big_string_for_word)



def replace_line_in_docx(file_path, target_string, target_header, stop_string, stop_header, replacement_string):
    sanitized_text = re.sub(r'[^\x20-\x7E]', '', replacement_string)

    doc = Document(file_path)

    flag = False
    
    for i, paragraph in enumerate(doc.paragraphs):
        # print(paragraph.text)
        # print()
        if paragraph.text.strip().split() == [target_string, target_header]:
            #print(paragraph.text)
            flag = True
            #new_paragraph = doc.add_paragraph(replacement_string)
            doc.paragraphs[i + 1].insert_paragraph_before(replacement_string)

        
        if paragraph.text.strip().split() == [stop_string, stop_header]:
            flag = False
            


        if flag:
            #p = paragraph._element
            #p.getparent().remove(p)
            print(paragraph.text)
        # if target_string in paragraph.text:
        #         doc.add_paragraph(sanitized_text)

    doc.save(file_path)

# Replace 'your_file.docx' with the actual path to your DOCX file. 
file_path = 'your_file.docx'
target_string = '2.4.2'
target_header = 'Hydrology'

stop_string = '2.4.3'
stop_header = 'Geology/Soils'

practice_text = "HELLO THIS IS WORKING "

replace_line_in_docx(copy_document_path, start_section, target_header, end_section, stop_header, new_text)




